import io

import docx


def extract_text_from_docx(docx_bytes: bytes) -> dict:
    """
    Extracts text from a .docx file's paragraphs (and table cells).

    Mirrors the return shape of extract_text_from_pdf so ingest_source's
    extract stage can treat both uniformly. .docx has no fixed pagination
    or bookmark TOC to build a section_map from, so both are None.

    Returns:
        {"text": str, "page_count": None, "section_map": None}
    """
    document = docx.Document(io.BytesIO(docx_bytes))

    parts = [p.text for p in document.paragraphs if p.text.strip()]
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    parts.append(cell.text)

    return {
        "text": "\n\n".join(parts),
        "page_count": None,
        "section_map": None,
    }
